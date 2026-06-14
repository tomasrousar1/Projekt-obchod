from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUTPUT = "Dokumentace_Projekt_obchod.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="BFC7D1", size="4"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths_cm):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for row in table.rows:
        for idx, width_cm in enumerate(widths_cm):
            cell = row.cells[idx]
            cell.width = Cm(width_cm)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_footer(section):
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Dokumentace projektu - Projekt Obchod")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(90, 90, 90)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    p.add_run(text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    p.add_run(text)
    return p


def add_label_paragraph(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(label)
    r.bold = True
    p.add_run(text)
    return p


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    add_footer(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor(31, 77, 120)
    title.paragraph_format.space_after = Pt(4)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(12)
    subtitle.font.color.rgb = RGBColor(90, 90, 90)
    subtitle.paragraph_format.space_after = Pt(14)

    h1 = styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(46, 116, 181)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)

    h2 = styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(46, 116, 181)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(7)

    for list_style in ("List Bullet", "List Number"):
        s = styles[list_style]
        s.font.name = "Calibri"
        s.font.size = Pt(11)
        s.paragraph_format.left_indent = Inches(0.375)
        s.paragraph_format.first_line_indent = Inches(-0.188)
        s.paragraph_format.space_after = Pt(4)
        s.paragraph_format.line_spacing = 1.25


def build_doc():
    doc = Document()
    style_document(doc)

    title = doc.add_paragraph(style="Title")
    title.add_run("Dokumentace projektu: Projekt Obchod")

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("Závěrečný projekt z programování v jazyce C#")

    p = doc.add_paragraph()
    p.add_run("Téma projektu: ").bold = True
    p.add_run("jednoduchý konzolový informační systém pro správu produktů v obchodě.")

    doc.add_heading("1. Zadání projektu", level=1)
    doc.add_paragraph(
        "Cílem projektu je vytvořit konzolovou aplikaci, která slouží jako jednoduchá evidence "
        "produktů ve skladu nebo obchodě. Program pracuje se seznamem produktů, umožňuje jejich "
        "výpis, filtrování podle kategorie, hromadné změny počtu kusů, přidání nového produktu "
        "a uložení aktuálního stavu do souboru."
    )
    add_label_paragraph(doc, "Použitý jazyk: ", "C# (.NET, konzolová aplikace).")
    add_label_paragraph(doc, "Hlavní princip: ", "práce s objekty třídy Produkty uloženými v dynamické kolekci List<Produkty>.")
    add_label_paragraph(doc, "Datový soubor: ", "Data.csv, ze kterého se produkty načítají a do kterého se ukládají změny.")

    doc.add_heading("2. Splnění podmínek projektu", level=1)
    add_bullet(doc, "Projekt je vytvořen v jazyce C# a používá objektově orientované programování.")
    add_bullet(doc, "Program obsahuje vlastní třídu Produkty, která popisuje jeden produkt v obchodě.")
    add_bullet(doc, "Produkty jsou uloženy v dynamické kolekci List<Produkty>, takže je možné přidávat další položky za běhu programu.")
    add_bullet(doc, "Aplikace načítá data ze souboru Data.csv pomocí StreamReader.")
    add_bullet(doc, "Aplikace ukládá data zpět do souboru Data.csv pomocí StreamWriter.")

    doc.add_heading("3. Model tříd a vazby", level=1)
    doc.add_paragraph(
        "Projekt je jednoduchý, proto obsahuje dvě hlavní části: třídu Program s hlavním během "
        "aplikace a třídu Produkty, která reprezentuje jednotlivé položky skladu."
    )

    table = doc.add_table(rows=1, cols=3)
    set_table_width(table, [3.2, 3.0, 10.3])
    set_table_borders(table)
    hdr = table.rows[0].cells
    hdr[0].text = "Třída"
    hdr[1].text = "Vazba"
    hdr[2].text = "Popis"
    for cell in hdr:
        set_cell_shading(cell, "E8EEF5")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    rows = [
        ("Program", "obsahuje List<Produkty>", "Řídí načítání dat, menu, zpracování voleb uživatele a ukládání dat."),
        ("Produkty", "datový objekt", "Uchovává informace o jednom produktu: ID, název, počet kusů, kategorii, cenu a minimální počet."),
        ("Data.csv", "zdroj dat", "Každý řádek souboru odpovídá jednomu objektu třídy Produkty."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    set_table_width(table, [3.2, 3.0, 10.3])

    doc.add_heading("4. Struktura aplikace", level=1)
    doc.add_heading("Třída Produkty", level=2)
    doc.add_paragraph("Třída Produkty slouží jako datový model. Obsahuje vlastnosti, které popisují jeden produkt v obchodě.")

    props = doc.add_table(rows=1, cols=3)
    set_table_width(props, [4.0, 3.0, 9.5])
    set_table_borders(props)
    header = props.rows[0].cells
    header[0].text = "Vlastnost"
    header[1].text = "Datový typ"
    header[2].text = "Význam"
    for cell in header:
        set_cell_shading(cell, "E8EEF5")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    properties = [
        ("ID", "int", "Jedinečné číslo produktu."),
        ("Name", "string", "Název produktu."),
        ("Pocet", "int", "Aktuální počet kusů na skladě."),
        ("Kategorie", "string", "Kategorie, do které produkt patří."),
        ("Cena", "decimal", "Cena produktu v Kč."),
        ("MinimalniPocet", "int", "Hranice, při které se produkt považuje za málo dostupný."),
    ]
    for prop in properties:
        cells = props.add_row().cells
        for idx, value in enumerate(prop):
            cells[idx].text = value
    set_table_width(props, [4.0, 3.0, 9.5])

    doc.add_heading("Třída Program a metoda Main", level=2)
    doc.add_paragraph(
        "Metoda Main obsahuje hlavní logiku programu. Na začátku vytvoří dynamický seznam produktů, "
        "načte data ze souboru a potom spustí nekonečný cyklus s uživatelským menu."
    )
    add_bullet(doc, "Vytvoření seznamu: List<Produkty> produtcs = new List<Produkty>().")
    add_bullet(doc, "Načtení CSV souboru pomocí StreamReader.")
    add_bullet(doc, "Převedení každého řádku souboru na objekt Produkty.")
    add_bullet(doc, "Zobrazení menu a zpracování volby uživatele pomocí switch.")
    add_bullet(doc, "Práce se seznamem pomocí cyklů foreach.")
    add_bullet(doc, "Uložení dat do souboru pomocí StreamWriter.")

    doc.add_heading("5. Popis práce se soubory", level=1)
    doc.add_paragraph(
        "Program používá soubor Data.csv. Tento soubor obsahuje tabulková data oddělená čárkami. "
        "První řádek obsahuje názvy sloupců a další řádky obsahují jednotlivé produkty."
    )
    add_label_paragraph(doc, "Struktura řádku: ", "ID,Name,Pocet,Kategorie,Cena,MinimalniPocet")
    add_label_paragraph(doc, "Příklad řádku: ", "1,Telefon,8,elektro,8990,3")

    doc.add_heading("Načítání dat", level=2)
    add_number(doc, "Program otevře soubor Data.csv pomocí StreamReader.")
    add_number(doc, "První řádek se přeskočí, protože obsahuje hlavičku.")
    add_number(doc, "Každý další řádek se rozdělí podle čárky pomocí Split(\",\").")
    add_number(doc, "Hodnoty se převedou na správné datové typy a uloží do nového objektu Produkty.")
    add_number(doc, "Objekt se přidá do seznamu produtcs.")

    doc.add_heading("Ukládání dat", level=2)
    doc.add_paragraph(
        "Při volbě 6 program otevře Data.csv pomocí StreamWriter a projde všechny produkty v seznamu. "
        "Každý produkt zapíše na samostatný řádek. Díky tomu se změny provedené během programu mohou uložit "
        "pro další spuštění aplikace."
    )

    doc.add_heading("6. Popis ovládání programu", level=1)
    doc.add_paragraph("Po spuštění se uživateli zobrazuje menu s šesti možnostmi:")
    controls = [
        ("1. Vypsat všechny produkty", "Zobrazí všechny produkty včetně ID, názvu, počtu, kategorie a ceny."),
        ("2. Vypsat produkty podle vybrané kategorie", "Uživatel zadá kategorii a program zobrazí pouze produkty z této kategorie."),
        ("3. Naskladnění nebo odebrání celé kategorie", "Uživatel vybere kategorii a rozhodne, zda chce počet kusů snížit nebo zvýšit."),
        ("4. Přidat nový produkt", "Program se zeptá na název, počet, kategorii a cenu. Nový produkt se přidá do seznamu."),
        ("5. Zobrazit produkty s nízkým počtem kusů", "Uživatel zadá maximální počet kusů a program vypíše produkty, které mají tento počet nebo méně."),
        ("6. Uložit data do souboru", "Aktuální seznam produktů se zapíše do souboru Data.csv."),
    ]
    for name, desc in controls:
        add_label_paragraph(doc, f"{name}: ", desc)

    doc.add_heading("7. Průběh programu", level=1)
    add_number(doc, "Program načte produkty ze souboru Data.csv.")
    add_number(doc, "Před zobrazením menu se u každého produktu sníží počet o 2 kusy, což simuluje úbytek zásob.")
    add_number(doc, "Pokud má některý produkt počet menší nebo roven minimálnímu počtu, program se zeptá, kolik kusů se má doplnit.")
    add_number(doc, "Uživatel vybere akci v menu.")
    add_number(doc, "Program provede vybranou akci a vrátí se zpět do menu.")
    add_number(doc, "Při volbě uložení se aktuální stav zapíše do CSV souboru.")

    doc.add_heading("8. Závěr", level=1)
    doc.add_paragraph(
        "Projekt splňuje zadání závěrečného projektu, protože je napsaný v C#, využívá objektově "
        "orientovaný přístup, pracuje s dynamickou kolekcí a načítá i ukládá data ze souboru. "
        "Aplikace představuje jednoduchý skladový systém, který je možné dále rozšiřovat například "
        "o úpravu jednotlivých produktů, mazání produktů nebo kontrolu chybně zadaných vstupů."
    )

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
